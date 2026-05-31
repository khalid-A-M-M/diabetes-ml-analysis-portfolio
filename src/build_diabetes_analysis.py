from __future__ import annotations

import json
import os
from pathlib import Path

PROJECT_DIR = Path(__file__).resolve().parents[1]
os.environ.setdefault("MPLCONFIGDIR", str(PROJECT_DIR / ".matplotlib"))
os.environ.setdefault("JOBLIB_TEMP_FOLDER", str(PROJECT_DIR / ".joblib_tmp"))

import matplotlib.pyplot as plt
import nbformat as nbf
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.shared import Inches
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import GradientBoostingClassifier, RandomForestClassifier
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    ConfusionMatrixDisplay,
    accuracy_score,
    confusion_matrix,
    f1_score,
    precision_score,
    recall_score,
    roc_auc_score,
    roc_curve,
)
from sklearn.model_selection import GridSearchCV, StratifiedKFold, train_test_split
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler


DATA_PATH = PROJECT_DIR / "data" / "diabetes.csv"
CHART_DIR = PROJECT_DIR / "assets" / "charts"
REPORT_DIR = PROJECT_DIR / "reports"
NOTEBOOK_DIR = PROJECT_DIR / "notebooks"

RANDOM_STATE = 42
MEDICAL_ZERO_AS_MISSING = ["Glucose", "BloodPressure", "SkinThickness", "Insulin", "BMI"]
FEATURES = [
    "Pregnancies",
    "Glucose",
    "BloodPressure",
    "SkinThickness",
    "Insulin",
    "BMI",
    "DiabetesPedigreeFunction",
    "Age",
]
TARGET = "Outcome"


def setup_style() -> None:
    sns.set_theme(style="whitegrid", context="notebook")
    plt.rcParams.update(
        {
            "figure.dpi": 130,
            "savefig.dpi": 180,
            "axes.titleweight": "bold",
            "axes.labelsize": 10,
            "axes.titlesize": 12,
            "font.size": 10,
        }
    )


def save_fig(filename: str) -> Path:
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    path = CHART_DIR / filename
    plt.tight_layout()
    plt.savefig(path, bbox_inches="tight")
    plt.close()
    return path


def df_to_markdown(df: pd.DataFrame, index: bool = True) -> str:
    table = df.copy()
    if index:
        table = table.reset_index()
    table = table.astype(str)
    headers = list(table.columns)
    rows = table.values.tolist()
    header_line = "| " + " | ".join(headers) + " |"
    separator = "| " + " | ".join(["---"] * len(headers)) + " |"
    body = ["| " + " | ".join(row) + " |" for row in rows]
    return "\n".join([header_line, separator, *body])


def load_data() -> pd.DataFrame:
    return pd.read_csv(DATA_PATH)


def assess_quality(df: pd.DataFrame) -> dict:
    zero_counts = df[MEDICAL_ZERO_AS_MISSING].eq(0).sum().to_dict()
    return {
        "records": int(df.shape[0]),
        "features": int(df.shape[1] - 1),
        "duplicate_rows": int(df.duplicated().sum()),
        "explicit_missing_values": df.isna().sum().astype(int).to_dict(),
        "zero_values_treated_as_missing": {k: int(v) for k, v in zero_counts.items()},
        "target_distribution": df[TARGET].value_counts().sort_index().astype(int).to_dict(),
    }


def prepare_data(df: pd.DataFrame) -> tuple[pd.DataFrame, dict]:
    clean = df.copy()
    clean[MEDICAL_ZERO_AS_MISSING] = clean[MEDICAL_ZERO_AS_MISSING].replace(0, np.nan)

    imputer = SimpleImputer(strategy="median")
    clean[MEDICAL_ZERO_AS_MISSING] = imputer.fit_transform(clean[MEDICAL_ZERO_AS_MISSING])

    cap_summary = {}
    for col in FEATURES:
        q1 = clean[col].quantile(0.25)
        q3 = clean[col].quantile(0.75)
        iqr = q3 - q1
        lower = q1 - 1.5 * iqr
        upper = q3 + 1.5 * iqr
        before = int(((clean[col] < lower) | (clean[col] > upper)).sum())
        clean[col] = clean[col].clip(lower, upper)
        cap_summary[col] = {"lower_cap": float(lower), "upper_cap": float(upper), "capped_values": before}

    clean["BMI_Category"] = pd.cut(
        clean["BMI"],
        bins=[0, 18.5, 25, 30, np.inf],
        labels=["Underweight", "Normal", "Overweight", "Obese"],
        include_lowest=True,
    )
    clean["Age_Group"] = pd.cut(
        clean["Age"],
        bins=[20, 30, 40, 50, 60, 100],
        labels=["21-30", "31-40", "41-50", "51-60", "60+"],
        include_lowest=True,
    )
    clean["Glucose_Risk_Group"] = pd.cut(
        clean["Glucose"],
        bins=[0, 99, 125, np.inf],
        labels=["Normal", "Elevated", "High"],
        include_lowest=True,
    )

    prep_summary = {
        "imputation_strategy": "Median imputation for medical measures where zero is not clinically valid.",
        "outlier_strategy": "IQR capping to reduce extreme-value influence while retaining records.",
        "outlier_caps": cap_summary,
        "engineered_features": ["BMI_Category", "Age_Group", "Glucose_Risk_Group"],
    }
    return clean, prep_summary


def run_eda(clean: pd.DataFrame) -> dict:
    desc = clean[FEATURES].agg(["mean", "median", "std", "min", "max"]).T.round(2)
    corr = clean[FEATURES + [TARGET]].corr(numeric_only=True)[TARGET].drop(TARGET).sort_values(ascending=False)

    outcome_summary = (
        clean.groupby(TARGET)[["Glucose", "BMI", "Age", "BloodPressure", "Insulin"]]
        .mean()
        .rename(index={0: "No Diabetes", 1: "Diabetes"})
        .round(2)
    )

    risk_by_group = {
        "BMI_Category": clean.groupby("BMI_Category", observed=True)[TARGET].mean().round(3).to_dict(),
        "Age_Group": clean.groupby("Age_Group", observed=True)[TARGET].mean().round(3).to_dict(),
        "Glucose_Risk_Group": clean.groupby("Glucose_Risk_Group", observed=True)[TARGET].mean().round(3).to_dict(),
    }

    return {
        "descriptive_statistics": desc,
        "target_correlations": corr.round(3),
        "outcome_summary": outcome_summary,
        "risk_by_group": risk_by_group,
    }


def build_visuals(clean: pd.DataFrame, model_outputs: dict | None = None) -> dict[str, str]:
    charts = {}

    fig, axes = plt.subplots(2, 2, figsize=(11, 8))
    for ax, col in zip(axes.ravel(), ["Glucose", "BMI", "Age", "BloodPressure"]):
        sns.histplot(data=clean, x=col, hue=TARGET, kde=True, bins=25, ax=ax, palette=["#2A9D8F", "#E76F51"])
        ax.set_title(f"{col} Distribution by Diabetes Outcome")
        ax.legend(title="Outcome", labels=["Diabetes", "No Diabetes"])
    charts["feature_distributions"] = str(save_fig("01_feature_distributions.png"))

    fig, axes = plt.subplots(2, 2, figsize=(11, 8))
    for ax, col in zip(axes.ravel(), ["Glucose", "BMI", "Insulin", "Age"]):
        sns.boxplot(data=clean, x=TARGET, y=col, ax=ax, palette=["#2A9D8F", "#E76F51"], hue=TARGET, legend=False)
        ax.set_xticks([0, 1])
        ax.set_xticklabels(["No Diabetes", "Diabetes"])
        ax.set_title(f"{col} by Outcome")
    charts["boxplots"] = str(save_fig("02_outcome_boxplots.png"))

    plt.figure(figsize=(9, 7))
    sns.heatmap(clean[FEATURES + [TARGET]].corr(), annot=True, cmap="RdBu_r", center=0, fmt=".2f", linewidths=0.5)
    plt.title("Correlation Heatmap: Diabetes Risk Indicators")
    charts["correlation_heatmap"] = str(save_fig("03_correlation_heatmap.png"))

    plt.figure(figsize=(8, 6))
    sns.scatterplot(
        data=clean,
        x="Glucose",
        y="BMI",
        hue=TARGET,
        size="Age",
        alpha=0.72,
        palette=["#2A9D8F", "#E76F51"],
    )
    plt.title("Glucose vs BMI, Sized by Age")
    plt.legend(title="Outcome / Age", bbox_to_anchor=(1.02, 1), loc="upper left")
    charts["scatter_glucose_bmi"] = str(save_fig("04_scatter_glucose_bmi.png"))

    pair = sns.pairplot(
        clean[["Glucose", "BMI", "Age", "BloodPressure", TARGET]],
        hue=TARGET,
        palette=["#2A9D8F", "#E76F51"],
        diag_kind="hist",
        corner=True,
        plot_kws={"alpha": 0.65, "s": 28},
    )
    pair.fig.suptitle("Pairplot of Key Diabetes Indicators", y=1.02, fontweight="bold")
    pair.savefig(CHART_DIR / "05_pairplot_key_indicators.png", bbox_inches="tight", dpi=180)
    plt.close("all")
    charts["pairplot"] = str(CHART_DIR / "05_pairplot_key_indicators.png")

    risk_data = clean.groupby("Glucose_Risk_Group", observed=True)[TARGET].mean().reset_index()
    plt.figure(figsize=(8, 5))
    sns.barplot(data=risk_data, x="Glucose_Risk_Group", y=TARGET, color="#E76F51")
    plt.title("Diabetes Rate by Glucose Risk Group")
    plt.ylabel("Diabetes Rate")
    plt.xlabel("Glucose Risk Group")
    charts["outcome_comparison"] = str(save_fig("06_diabetes_rate_by_glucose_group.png"))

    if model_outputs:
        metrics_df = model_outputs["metrics"].sort_values("ROC_AUC", ascending=False)
        plt.figure(figsize=(9, 5))
        metrics_df.set_index("Model")[["Accuracy", "Precision", "Recall", "F1", "ROC_AUC"]].plot(kind="bar", ax=plt.gca())
        plt.title("Model Performance Comparison")
        plt.ylabel("Score")
        plt.ylim(0, 1)
        plt.xticks(rotation=20, ha="right")
        charts["model_comparison"] = str(save_fig("07_model_performance_comparison.png"))

        best_model_name = model_outputs["best_model_name"]
        cm = model_outputs["confusion_matrices"][best_model_name]
        disp = ConfusionMatrixDisplay(confusion_matrix=cm, display_labels=["No Diabetes", "Diabetes"])
        disp.plot(cmap="Blues", values_format="d")
        plt.title(f"Confusion Matrix: {best_model_name}")
        charts["confusion_matrix"] = str(save_fig("08_best_model_confusion_matrix.png"))

        plt.figure(figsize=(7, 5))
        for model_name, roc_data in model_outputs["roc_curves"].items():
            plt.plot(roc_data["fpr"], roc_data["tpr"], label=f"{model_name} AUC={roc_data['auc']:.3f}")
        plt.plot([0, 1], [0, 1], linestyle="--", color="gray")
        plt.title("ROC Curves by Model")
        plt.xlabel("False Positive Rate")
        plt.ylabel("True Positive Rate")
        plt.legend()
        charts["roc_curves"] = str(save_fig("09_roc_curves.png"))

        importance = model_outputs["feature_importance"].sort_values("Importance", ascending=True)
        plt.figure(figsize=(8, 5))
        sns.barplot(data=importance, x="Importance", y="Feature", color="#264653")
        plt.title(f"Feature Importance: {best_model_name}")
        charts["feature_importance"] = str(save_fig("10_feature_importance.png"))

    return charts


def train_models(clean: pd.DataFrame) -> dict:
    X = clean[FEATURES]
    y = clean[TARGET]
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=RANDOM_STATE, stratify=y
    )

    cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=RANDOM_STATE)
    scaler_preprocessor = ColumnTransformer(
        transformers=[("numeric", Pipeline([("imputer", SimpleImputer(strategy="median")), ("scaler", StandardScaler())]), FEATURES)]
    )
    tree_preprocessor = ColumnTransformer(
        transformers=[("numeric", Pipeline([("imputer", SimpleImputer(strategy="median"))]), FEATURES)]
    )

    model_specs = {
        "Logistic Regression": (
            Pipeline([("prep", scaler_preprocessor), ("model", LogisticRegression(max_iter=1000, random_state=RANDOM_STATE))]),
            {"model__C": [0.1, 1.0, 3.0]},
        ),
        "Random Forest": (
            Pipeline([("prep", tree_preprocessor), ("model", RandomForestClassifier(random_state=RANDOM_STATE))]),
            {"model__n_estimators": [150, 250], "model__max_depth": [3, 5, None], "model__min_samples_leaf": [2, 5]},
        ),
        "Gradient Boosting": (
            Pipeline([("prep", tree_preprocessor), ("model", GradientBoostingClassifier(random_state=RANDOM_STATE))]),
            {"model__n_estimators": [80, 120], "model__learning_rate": [0.03, 0.06, 0.1], "model__max_depth": [2, 3]},
        ),
    }

    fitted_models = {}
    rows = []
    confusion_matrices = {}
    roc_curves = {}

    for name, (pipeline, grid) in model_specs.items():
        search = GridSearchCV(pipeline, grid, scoring="roc_auc", cv=cv, n_jobs=1)
        search.fit(X_train, y_train)
        best = search.best_estimator_
        y_pred = best.predict(X_test)
        y_proba = best.predict_proba(X_test)[:, 1]
        auc = roc_auc_score(y_test, y_proba)
        fpr, tpr, _ = roc_curve(y_test, y_proba)

        fitted_models[name] = best
        confusion_matrices[name] = confusion_matrix(y_test, y_pred)
        roc_curves[name] = {"fpr": fpr, "tpr": tpr, "auc": auc}
        rows.append(
            {
                "Model": name,
                "Best_Params": json.dumps(search.best_params_),
                "CV_ROC_AUC": round(float(search.best_score_), 3),
                "Accuracy": round(float(accuracy_score(y_test, y_pred)), 3),
                "Precision": round(float(precision_score(y_test, y_pred)), 3),
                "Recall": round(float(recall_score(y_test, y_pred)), 3),
                "F1": round(float(f1_score(y_test, y_pred)), 3),
                "ROC_AUC": round(float(auc), 3),
            }
        )

    metrics = pd.DataFrame(rows).sort_values("ROC_AUC", ascending=False)
    best_model_name = metrics.iloc[0]["Model"]
    best_model = fitted_models[best_model_name]
    model_step = best_model.named_steps["model"]

    if hasattr(model_step, "feature_importances_"):
        importance_values = model_step.feature_importances_
    else:
        importance_values = np.abs(model_step.coef_[0])
    feature_importance = pd.DataFrame({"Feature": FEATURES, "Importance": importance_values})
    feature_importance["Importance"] = feature_importance["Importance"] / feature_importance["Importance"].sum()

    return {
        "metrics": metrics,
        "best_model_name": str(best_model_name),
        "confusion_matrices": confusion_matrices,
        "roc_curves": roc_curves,
        "feature_importance": feature_importance.sort_values("Importance", ascending=False),
        "test_size": int(len(y_test)),
        "train_size": int(len(y_train)),
    }


def write_source_notes(quality: dict) -> None:
    source = PROJECT_DIR / "DATASET_SOURCE.md"
    source.write_text(
        "\n".join(
            [
                "# Dataset Source",
                "",
                "Dataset: Pima Indians Diabetes Database",
                "Downloaded CSV: https://raw.githubusercontent.com/plotly/datasets/master/diabetes.csv",
                "Original dataset reference: National Institute of Diabetes and Digestive and Kidney Diseases; widely mirrored through UCI/Kaggle learning repositories.",
                "",
                f"Records: {quality['records']}",
                f"Predictor features: {quality['features']}",
                "Target variable: Outcome, where 1 indicates diabetes and 0 indicates no diabetes.",
                "",
                "Important note: In this dataset, zeros in Glucose, BloodPressure, SkinThickness, Insulin, and BMI represent invalid physiological values and were treated as missing data during preprocessing.",
            ]
        ),
        encoding="utf-8",
    )


def write_report_md(quality: dict, prep: dict, eda: dict, model_outputs: dict, charts: dict) -> Path:
    metrics_md = df_to_markdown(model_outputs["metrics"], index=False)
    desc_md = df_to_markdown(eda["descriptive_statistics"])
    outcome_md = df_to_markdown(eda["outcome_summary"])
    corr_md = df_to_markdown(eda["target_correlations"].to_frame("Correlation with Outcome"))
    importance_md = df_to_markdown(model_outputs["feature_importance"].round(3), index=False)

    report = f"""# Diabetes Risk Analysis and Prediction Report

## Executive Summary

This project analyzes a real public diabetes dataset to identify the health indicators most associated with diabetes and build a predictive model that can estimate diabetes risk from patient characteristics.

The analysis found that Glucose is the strongest risk signal, followed by BMI and Age. Patients with higher glucose values and higher BMI showed visibly higher diabetes rates. The best-performing model was **{model_outputs['best_model_name']}**, selected using ROC-AUC on the test set.

## Dataset Overview

- Source: Pima Indians Diabetes Database, mirrored as a public CSV by Plotly datasets.
- Records: {quality['records']}
- Predictor features: {quality['features']}
- Target: `Outcome` where 1 = diabetes and 0 = no diabetes.
- Class distribution: {quality['target_distribution']}

## Business Problem

Healthcare teams often need to identify which patient indicators deserve the most attention during screening. This project answers three practical questions:

1. Which biometric variables are most associated with diabetes?
2. Can diabetes risk be predicted from standard patient measurements?
3. Which indicators should be highlighted in a business dashboard for health monitoring?

## Methodology

### Data Cleaning and Preparation

- Checked missing values, data types, duplicate rows, and target distribution.
- Found no explicit missing values, but detected physiologically invalid zeros in selected medical fields.
- Treated zeros in {", ".join(MEDICAL_ZERO_AS_MISSING)} as missing values.
- Applied median imputation to preserve all records.
- Used IQR capping for outliers to reduce extreme-value influence.
- Engineered interpretable features: BMI category, age group, and glucose risk group.

Duplicate records detected: {quality['duplicate_rows']}

Zero values treated as missing:
{json.dumps(quality['zero_values_treated_as_missing'], indent=2)}

### Exploratory Analysis

Descriptive statistics:

{desc_md}

Average values by outcome:

{outcome_md}

Correlation with diabetes outcome:

{corr_md}

## Visual Findings

Recommended dashboard charts:

- Feature distributions by outcome: `{Path(charts['feature_distributions']).name}`
- Outcome boxplots: `{Path(charts['boxplots']).name}`
- Correlation heatmap: `{Path(charts['correlation_heatmap']).name}`
- Glucose vs BMI scatterplot: `{Path(charts['scatter_glucose_bmi']).name}`
- Pairplot of core indicators: `{Path(charts['pairplot']).name}`
- Model comparison and feature importance charts.

## Predictive Modeling

Models tested:

- Logistic Regression
- Random Forest
- Gradient Boosting, used as a practical boosting alternative because XGBoost is optional and may not be available in all client environments.

Performance comparison:

{metrics_md}

Best model: **{model_outputs['best_model_name']}**

## Explainability

Feature importance:

{importance_md}

The strongest indicators should be interpreted as screening signals rather than clinical diagnosis. The model is useful for prioritization, reporting, and risk segmentation, but final diagnosis requires medical assessment.

## Business Insights

- Glucose is the clearest driver of diabetes risk in this dataset.
- BMI adds an important metabolic-risk signal and should be monitored alongside glucose.
- Age is associated with increased risk and helps segment patients into prevention groups.
- Blood Pressure and Insulin provide supporting context but are less predictive than glucose in this analysis.

## Recommendations

- Build a screening dashboard with KPI cards for diabetes rate, average glucose, average BMI, and high-risk patient count.
- Segment patients by glucose risk group, BMI category, and age group.
- Use the predictive model as a prioritization tool to flag records for follow-up.
- Communicate model outputs as risk estimates, not as a medical diagnosis.
"""
    path = REPORT_DIR / "Diabetes_Risk_Analysis_Report.md"
    path.write_text(report, encoding="utf-8")
    return path


def write_dashboard_spec() -> Path:
    spec = """# Power BI Dashboard Design Specification

## Page 1: Executive Overview

KPI Cards:
- Total Patients
- Diabetes Rate
- Average Glucose
- Average BMI
- High Glucose Patient Count

Filters:
- Outcome
- Age Group
- BMI Category
- Glucose Risk Group

Visual Components:
- Diabetes rate by glucose risk group
- Diabetes rate by BMI category
- Age group comparison
- Key indicator distribution panel

## Page 2: Risk Factor Analysis

Visual Components:
- Correlation heatmap image or matrix visual
- Boxplots for Glucose, BMI, Age, BloodPressure by Outcome
- Scatterplot: Glucose vs BMI, colored by Outcome and sized by Age
- Top feature importance bar chart

## Page 3: Model Performance

Visual Components:
- Model comparison table
- ROC-AUC comparison bar chart
- Confusion matrix for the selected model
- Feature importance ranking

## Suggested Layout

Use a clean healthcare analytics layout with a white background, muted teal for non-diabetes, warm coral for diabetes, and dark navy for headings. Keep KPI cards at the top, filters on the left, and analytical visuals in a two-column grid.
"""
    path = REPORT_DIR / "PowerBI_Dashboard_Design_Specification.md"
    path.write_text(spec, encoding="utf-8")
    return path


def write_docx_report(markdown_report_path: Path, charts: dict) -> Path:
    doc = Document()
    doc.add_heading("Diabetes Risk Analysis and Predictive Modeling", level=0)
    doc.add_paragraph("Professional report for a diabetes risk analysis and prediction project.")

    sections = [
        ("Executive Summary", "This project identifies the strongest indicators associated with diabetes and compares machine learning models for risk prediction."),
        ("Methodology", "The workflow includes data quality assessment, invalid-zero treatment, median imputation, outlier capping, exploratory analysis, visualization, model training, and model evaluation."),
        ("Key Findings", "Glucose is the strongest diabetes risk indicator in the dataset. BMI and age also show meaningful association with diabetes outcome."),
        ("Recommendations", "Healthcare teams should monitor glucose, BMI, and age-group risk segments, and use predictive models as prioritization tools rather than diagnostic replacements."),
    ]
    for title, body in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(body)

    doc.add_heading("Selected Visualizations", level=1)
    for key in ["correlation_heatmap", "scatter_glucose_bmi", "model_comparison", "feature_importance"]:
        chart = Path(charts[key])
        doc.add_paragraph(chart.stem.replace("_", " ").title())
        doc.add_picture(str(chart), width=Inches(6.2))

    doc.add_heading("Full Markdown Report", level=1)
    doc.add_paragraph(f"See: {markdown_report_path.name}")

    path = REPORT_DIR / "Diabetes_Risk_Analysis_Report.docx"
    doc.save(path)
    return path


def write_notebook() -> Path:
    nb = nbf.v4.new_notebook()
    cells = []

    cells.append(nbf.v4.new_markdown_cell("# Diabetes Risk Analysis and Predictive Modeling\n\nNotebook covering data cleaning, EDA, visualization, predictive modeling, explainability, and healthcare analytics recommendations."))
    cells.append(nbf.v4.new_markdown_cell("## 1. Business Understanding\n\nObjective: identify which patient characteristics are most associated with diabetes and build a practical model to estimate risk from biometric indicators."))
    cells.append(nbf.v4.new_code_cell("""from pathlib import Path
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns

from sklearn.model_selection import train_test_split, StratifiedKFold, GridSearchCV
from sklearn.pipeline import Pipeline
from sklearn.compose import ColumnTransformer
from sklearn.preprocessing import StandardScaler
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, roc_auc_score, confusion_matrix, roc_curve, ConfusionMatrixDisplay

sns.set_theme(style="whitegrid", context="notebook")
PROJECT_DIR = Path("..").resolve()
DATA_PATH = PROJECT_DIR / "data" / "diabetes.csv"
CHART_DIR = PROJECT_DIR / "assets" / "charts"
CHART_DIR.mkdir(parents=True, exist_ok=True)
RANDOM_STATE = 42
FEATURES = ["Pregnancies", "Glucose", "BloodPressure", "SkinThickness", "Insulin", "BMI", "DiabetesPedigreeFunction", "Age"]
TARGET = "Outcome"
MEDICAL_ZERO_AS_MISSING = ["Glucose", "BloodPressure", "SkinThickness", "Insulin", "BMI"]"""))

    cells.append(nbf.v4.new_markdown_cell("## 2. Dataset Overview\n\nThe dataset is the Pima Indians Diabetes Database. It contains 768 patient records, 8 predictor features, and a binary target `Outcome`."))
    cells.append(nbf.v4.new_code_cell("""df = pd.read_csv(DATA_PATH)
print(df.shape)
df.head()"""))
    cells.append(nbf.v4.new_code_cell("""quality = {
    "duplicates": df.duplicated().sum(),
    "missing_values": df.isna().sum(),
    "invalid_zeros": df[MEDICAL_ZERO_AS_MISSING].eq(0).sum(),
    "target_distribution": df[TARGET].value_counts().sort_index()
}
quality"""))

    cells.append(nbf.v4.new_markdown_cell("## 3. Data Cleaning and Preparation\n\nZero values in Glucose, BloodPressure, SkinThickness, Insulin, and BMI are not clinically valid, so they are treated as missing values and imputed with the median. Outliers are capped using the IQR rule."))
    cells.append(nbf.v4.new_code_cell("""clean = df.copy()
clean[MEDICAL_ZERO_AS_MISSING] = clean[MEDICAL_ZERO_AS_MISSING].replace(0, np.nan)
clean[MEDICAL_ZERO_AS_MISSING] = SimpleImputer(strategy="median").fit_transform(clean[MEDICAL_ZERO_AS_MISSING])

outlier_summary = {}
for col in FEATURES:
    q1, q3 = clean[col].quantile([0.25, 0.75])
    iqr = q3 - q1
    lower, upper = q1 - 1.5 * iqr, q3 + 1.5 * iqr
    outlier_summary[col] = int(((clean[col] < lower) | (clean[col] > upper)).sum())
    clean[col] = clean[col].clip(lower, upper)

clean["BMI_Category"] = pd.cut(clean["BMI"], [0, 18.5, 25, 30, np.inf], labels=["Underweight", "Normal", "Overweight", "Obese"], include_lowest=True)
clean["Age_Group"] = pd.cut(clean["Age"], [20, 30, 40, 50, 60, 100], labels=["21-30", "31-40", "41-50", "51-60", "60+"], include_lowest=True)
clean["Glucose_Risk_Group"] = pd.cut(clean["Glucose"], [0, 99, 125, np.inf], labels=["Normal", "Elevated", "High"], include_lowest=True)
outlier_summary"""))

    cells.append(nbf.v4.new_markdown_cell("## 4. Exploratory Data Analysis"))
    cells.append(nbf.v4.new_code_cell("""clean[FEATURES].agg(["mean", "median", "std", "min", "max"]).T.round(2)"""))
    cells.append(nbf.v4.new_code_cell("""clean.groupby(TARGET)[["Glucose", "BMI", "Age", "BloodPressure", "Insulin"]].mean().rename(index={0: "No Diabetes", 1: "Diabetes"}).round(2)"""))
    cells.append(nbf.v4.new_code_cell("""clean[FEATURES + [TARGET]].corr(numeric_only=True)[TARGET].drop(TARGET).sort_values(ascending=False).round(3)"""))

    cells.append(nbf.v4.new_markdown_cell("## 5. Professional Visualizations"))
    cells.append(nbf.v4.new_code_cell("""fig, axes = plt.subplots(2, 2, figsize=(11, 8))
for ax, col in zip(axes.ravel(), ["Glucose", "BMI", "Age", "BloodPressure"]):
    sns.histplot(data=clean, x=col, hue=TARGET, kde=True, bins=25, ax=ax, palette=["#2A9D8F", "#E76F51"])
    ax.set_title(f"{col} Distribution by Diabetes Outcome")
plt.tight_layout()"""))
    cells.append(nbf.v4.new_code_cell("""plt.figure(figsize=(9, 7))
sns.heatmap(clean[FEATURES + [TARGET]].corr(), annot=True, cmap="RdBu_r", center=0, fmt=".2f")
plt.title("Correlation Heatmap: Diabetes Risk Indicators")
plt.tight_layout()"""))
    cells.append(nbf.v4.new_code_cell("""plt.figure(figsize=(8, 6))
sns.scatterplot(data=clean, x="Glucose", y="BMI", hue=TARGET, size="Age", alpha=0.72, palette=["#2A9D8F", "#E76F51"])
plt.title("Glucose vs BMI, Sized by Age")
plt.tight_layout()"""))
    cells.append(nbf.v4.new_code_cell("""sns.pairplot(clean[["Glucose", "BMI", "Age", "BloodPressure", TARGET]], hue=TARGET, palette=["#2A9D8F", "#E76F51"], diag_kind="hist", corner=True)"""))

    cells.append(nbf.v4.new_markdown_cell("## 6. Predictive Modeling\n\nThree models are compared: Logistic Regression, Random Forest, and Gradient Boosting. XGBoost can be added in a production environment if available, but Gradient Boosting is used here as a reliable boosting baseline."))
    cells.append(nbf.v4.new_code_cell("""X = clean[FEATURES]
y = clean[TARGET]
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=RANDOM_STATE, stratify=y)

cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=RANDOM_STATE)
scaler_preprocessor = ColumnTransformer([("numeric", Pipeline([("imputer", SimpleImputer(strategy="median")), ("scaler", StandardScaler())]), FEATURES)])
tree_preprocessor = ColumnTransformer([("numeric", Pipeline([("imputer", SimpleImputer(strategy="median"))]), FEATURES)])

model_specs = {
    "Logistic Regression": (Pipeline([("prep", scaler_preprocessor), ("model", LogisticRegression(max_iter=1000, random_state=RANDOM_STATE))]), {"model__C": [0.1, 1.0, 3.0]}),
    "Random Forest": (Pipeline([("prep", tree_preprocessor), ("model", RandomForestClassifier(random_state=RANDOM_STATE))]), {"model__n_estimators": [150, 250], "model__max_depth": [3, 5, None], "model__min_samples_leaf": [2, 5]}),
    "Gradient Boosting": (Pipeline([("prep", tree_preprocessor), ("model", GradientBoostingClassifier(random_state=RANDOM_STATE))]), {"model__n_estimators": [80, 120], "model__learning_rate": [0.03, 0.06, 0.1], "model__max_depth": [2, 3]}),
}

fitted_models, rows, roc_data = {}, [], {}
for name, (pipeline, grid) in model_specs.items():
    search = GridSearchCV(pipeline, grid, scoring="roc_auc", cv=cv, n_jobs=1)
    search.fit(X_train, y_train)
    best = search.best_estimator_
    y_pred = best.predict(X_test)
    y_proba = best.predict_proba(X_test)[:, 1]
    fitted_models[name] = best
    fpr, tpr, _ = roc_curve(y_test, y_proba)
    roc_data[name] = (fpr, tpr, roc_auc_score(y_test, y_proba))
    rows.append({
        "Model": name,
        "CV_ROC_AUC": search.best_score_,
        "Accuracy": accuracy_score(y_test, y_pred),
        "Precision": precision_score(y_test, y_pred),
        "Recall": recall_score(y_test, y_pred),
        "F1": f1_score(y_test, y_pred),
        "ROC_AUC": roc_auc_score(y_test, y_proba),
        "Best_Params": search.best_params_
    })

metrics = pd.DataFrame(rows).sort_values("ROC_AUC", ascending=False)
metrics.round(3)"""))

    cells.append(nbf.v4.new_markdown_cell("## 7. Explainability"))
    cells.append(nbf.v4.new_code_cell("""best_model_name = metrics.iloc[0]["Model"]
best_model = fitted_models[best_model_name]
model_step = best_model.named_steps["model"]
if hasattr(model_step, "feature_importances_"):
    importance_values = model_step.feature_importances_
else:
    importance_values = np.abs(model_step.coef_[0])

feature_importance = pd.DataFrame({"Feature": FEATURES, "Importance": importance_values})
feature_importance["Importance"] = feature_importance["Importance"] / feature_importance["Importance"].sum()
feature_importance.sort_values("Importance", ascending=False).round(3)"""))
    cells.append(nbf.v4.new_code_cell("""plt.figure(figsize=(8, 5))
sns.barplot(data=feature_importance.sort_values("Importance"), x="Importance", y="Feature", color="#264653")
plt.title(f"Feature Importance: {best_model_name}")
plt.tight_layout()"""))

    cells.append(nbf.v4.new_markdown_cell("## 8. Business Insights\n\n- Glucose is the clearest diabetes risk signal.\n- BMI and Age provide practical segmentation dimensions for screening and dashboards.\n- Predictive modeling can prioritize follow-up, but model results should not be framed as clinical diagnosis.\n- Recommended dashboard KPIs: diabetes rate, average glucose, average BMI, high-risk glucose count, and patient count by age/BMI segment."))

    nb["cells"] = cells
    nb["metadata"] = {
        "kernelspec": {"display_name": "Python 3", "language": "python", "name": "python3"},
        "language_info": {"name": "python", "version": "3.12"},
    }

    path = NOTEBOOK_DIR / "Diabetes_Risk_Analysis.ipynb"
    nbf.write(nb, path)
    return path


def write_readme(model_outputs: dict) -> Path:
    readme = f"""# Diabetes Risk Analysis and Prediction

## Overview

This project analyzes a real public diabetes dataset and builds Machine Learning models to predict diabetes risk from patient health indicators. The goal is to identify the strongest risk factors, compare model performance, and present the results in a clear format that supports healthcare screening and decision-making.

## Dataset

The analysis uses the **Pima Indians Diabetes Database**, a public medical dataset with 768 records.

The dataset includes:

- Pregnancies
- Glucose
- Blood pressure
- Skin thickness
- Insulin
- BMI
- Diabetes pedigree function
- Age
- Diabetes outcome

Dataset source details are documented in [DATASET_SOURCE.md](DATASET_SOURCE.md).

## Business Questions

This project answers practical healthcare analytics questions:

- Which health indicators are most associated with diabetes?
- How different are glucose, BMI, and age patterns between diabetic and non-diabetic patients?
- Which Machine Learning model performs best for diabetes risk prediction?
- Which features should be highlighted in a health monitoring dashboard?

## Work Completed

1. Loaded and validated the public diabetes dataset.
2. Checked missing values, duplicate records, invalid medical zeros, and target distribution.
3. Treated invalid zero values in medical fields as missing data and applied median imputation.
4. Created interpretable features such as BMI category, age group, and glucose risk group.
5. Built exploratory charts for distributions, outcome comparisons, correlations, and risk groups.
6. Trained and compared multiple Machine Learning models.
7. Selected {model_outputs['best_model_name']} as the best-performing model based on ROC-AUC.
8. Generated model performance tables, feature importance, charts, and a final analysis report.

## Key Findings

- Glucose is the strongest predictor of diabetes risk.
- BMI and age also show meaningful association with diabetes outcome.
- Patients with higher glucose and higher BMI show visibly higher diabetes rates.
- {model_outputs['best_model_name']} achieved the strongest overall model performance among the tested models.

## Model Summary

Best model: **{model_outputs['best_model_name']}**

Model outputs are available in:

- [model_performance.csv](reports/model_performance.csv)
- [feature_importance.csv](reports/feature_importance.csv)
- [Diabetes_Risk_Analysis_Report.md](reports/Diabetes_Risk_Analysis_Report.md)

## Visual Results

### Feature Distributions

![Feature Distributions](assets/charts/01_feature_distributions.png)

### Outcome Boxplots

![Outcome Boxplots](assets/charts/02_outcome_boxplots.png)

### Correlation Heatmap

![Correlation Heatmap](assets/charts/03_correlation_heatmap.png)

### Model Performance

![Model Performance](assets/charts/07_model_performance_comparison.png)

### Confusion Matrix

![Confusion Matrix](assets/charts/08_best_model_confusion_matrix.png)

### Feature Importance

![Feature Importance](assets/charts/10_feature_importance.png)

## Project Files

```text
assets/charts/                         # Generated visualizations
data/                                  # Dataset files
notebooks/Diabetes_Risk_Analysis.ipynb # Analysis notebook
reports/                               # Report, metrics, and model outputs
src/build_diabetes_analysis.py         # Reproducible analysis pipeline
DATASET_SOURCE.md                      # Dataset documentation
```

## How to Run

Run the full analysis pipeline:

```bash
python src/build_diabetes_analysis.py
```

The script regenerates the analysis outputs, charts, model metrics, report, and notebook.

## Tools Used

- Python
- Pandas
- NumPy
- Matplotlib
- Seaborn
- Scikit-learn
- Random Forest Classification

## Deliverables

- Cleaned and prepared medical dataset
- Exploratory data analysis charts
- Model comparison and evaluation metrics
- Feature importance analysis
- Final diabetes risk analysis report
- Reproducible Python pipeline
"""
    path = PROJECT_DIR / "README.md"
    path.write_text(readme, encoding="utf-8")
    return path


def main() -> None:
    setup_style()
    for directory in [CHART_DIR, REPORT_DIR, NOTEBOOK_DIR]:
        directory.mkdir(parents=True, exist_ok=True)

    df = load_data()
    quality = assess_quality(df)
    clean, prep = prepare_data(df)
    eda = run_eda(clean)
    preliminary_charts = build_visuals(clean)
    model_outputs = train_models(clean)
    charts = build_visuals(clean, model_outputs)

    clean.to_csv(PROJECT_DIR / "data" / "diabetes_cleaned.csv", index=False)
    model_outputs["metrics"].to_csv(REPORT_DIR / "model_performance.csv", index=False)
    model_outputs["feature_importance"].to_csv(REPORT_DIR / "feature_importance.csv", index=False)
    eda["descriptive_statistics"].to_csv(REPORT_DIR / "descriptive_statistics.csv")

    write_source_notes(quality)
    report_md = write_report_md(quality, prep, eda, model_outputs, charts)
    write_dashboard_spec()
    write_docx_report(report_md, charts)
    write_notebook()
    write_readme(model_outputs)

    summary = {
        "project_dir": str(PROJECT_DIR),
        "records": quality["records"],
        "features": quality["features"],
        "best_model": model_outputs["best_model_name"],
        "best_model_metrics": model_outputs["metrics"].iloc[0].to_dict(),
        "charts_created": len(charts),
    }
    (REPORT_DIR / "project_summary.json").write_text(json.dumps(summary, indent=2), encoding="utf-8")
    print(json.dumps(summary, indent=2))


if __name__ == "__main__":
    main()
